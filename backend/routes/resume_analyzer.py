from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
from models.resume import ResumeData, ResumeAnalysis, ResumeImprovement
from typing import List
import PyPDF2
import io
import json
import docx
import requests
import os
from dotenv import load_dotenv

load_dotenv()

router = APIRouter()

# Gemini API Configuration
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent"

# Response Model
class AnalyzeResponse(BaseModel):
    overall_score: int
    content_quality: int
    formatting_score: int
    keyword_match: int
    experience_impact: int
    skills_relevance: int
    ats_compatibility: int
    strengths: List[str]
    improvements: List[str]
    missing_keywords: List[str]
    suggested_skills: List[str]
    extracted_text: str

@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_resume(file: UploadFile = File(...)):
    """Analyze uploaded resume (PDF/DOC/DOCX) and return AI-powered insights"""
    
    try:
        # Extract text based on file type
        extracted_text = ""
        
        if file.filename.endswith('.pdf'):
            # Read PDF
            content = await file.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
                
        elif file.filename.endswith('.docx'):
            # Read DOCX
            content = await file.read()
            doc = docx.Document(io.BytesIO(content))
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
                
        elif file.filename.endswith('.doc'):
            raise HTTPException(status_code=400, detail=".doc format not supported. Please use .docx or .pdf")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF or DOCX")
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the document")
        
        # Use Gemini AI for dynamic analysis via REST API
        if GEMINI_API_KEY:
            try:
                prompt = f"""You are an expert resume analyzer and career coach. Analyze the following resume text and provide detailed feedback.

Resume Text:
{extracted_text[:3000]}

Please provide a comprehensive analysis in the following JSON format:
{{
    "overall_score": <number 0-100>,
    "content_quality": <number 0-100>,
    "formatting_score": <number 0-100>,
    "keyword_match": <number 0-100>,
    "experience_impact": <number 0-100>,
    "skills_relevance": <number 0-100>,
    "ats_compatibility": <number 0-100>,
    "strengths": [<list of 4-6 specific strengths>],
    "improvements": [<list of 4-6 actionable improvements>],
    "missing_keywords": [<list of 6-8 important missing keywords/skills>],
    "suggested_skills": [<list of 5-7 skills to add based on career path>]
}}

Scoring criteria:
- overall_score: Overall resume quality
- content_quality: Quality of written content, grammar, clarity
- formatting_score: Structure, readability, visual organization
- keyword_match: Presence of industry-relevant keywords
- experience_impact: How well experience demonstrates impact and achievements
- skills_relevance: Relevance and demand of listed skills
- ats_compatibility: How well it would parse through ATS systems

Return ONLY valid JSON, no markdown or explanations."""

                # Call Gemini API directly
                headers = {"Content-Type": "application/json"}
                payload = {
                    "contents": [{
                        "parts": [{"text": prompt}]
                    }]
                }
                
                response = requests.post(
                    f"{GEMINI_API_URL}?key={GEMINI_API_KEY}",
                    headers=headers,
                    json=payload,
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    ai_text = result.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
                    
                    # Clean the response (remove markdown code blocks if present)
                    if "```json" in ai_text:
                        ai_text = ai_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in ai_text:
                        ai_text = ai_text.split("```")[1].split("```")[0].strip()
                    
                    # Parse AI response
                    ai_analysis = json.loads(ai_text)
                    
                    return AnalyzeResponse(
                        overall_score=int(ai_analysis.get("overall_score", 85)),
                        content_quality=int(ai_analysis.get("content_quality", 85)),
                        formatting_score=int(ai_analysis.get("formatting_score", 82)),
                        keyword_match=int(ai_analysis.get("keyword_match", 80)),
                        experience_impact=int(ai_analysis.get("experience_impact", 83)),
                        skills_relevance=int(ai_analysis.get("skills_relevance", 85)),
                        ats_compatibility=int(ai_analysis.get("ats_compatibility", 80)),
                        strengths=ai_analysis.get("strengths", [])[:6],
                        improvements=ai_analysis.get("improvements", [])[:6],
                        missing_keywords=ai_analysis.get("missing_keywords", [])[:8],
                        suggested_skills=ai_analysis.get("suggested_skills", [])[:7],
                        extracted_text=extracted_text[:500]
                    )
                else:
                    print(f"Gemini API Error: {response.status_code} - {response.text}")
                    return await fallback_analysis(extracted_text)
                    
            except json.JSONDecodeError as je:
                print(f"JSON Parse Error: {je}")
                return await fallback_analysis(extracted_text)
            except Exception as ai_error:
                print(f"Gemini API Error: {ai_error}")
                return await fallback_analysis(extracted_text)
        else:
            return await fallback_analysis(extracted_text)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing resume: {str(e)}")


async def fallback_analysis(extracted_text: str) -> AnalyzeResponse:
    """Fallback analysis method when AI fails - returns optimistic scores"""
    text_lower = extracted_text.lower()
    
    # Calculate scores based on content analysis
    has_education = any(keyword in text_lower for keyword in ['education', 'degree', 'university', 'college', 'bachelor', 'master'])
    has_experience = any(keyword in text_lower for keyword in ['experience', 'worked', 'developed', 'managed', 'led'])
    has_skills = any(keyword in text_lower for keyword in ['skills', 'python', 'java', 'javascript', 'react', 'node'])
    has_contact = any(keyword in text_lower for keyword in ['email', 'phone', 'linkedin', '@'])
    
    word_count = len(extracted_text.split())
    line_count = len(extracted_text.split('\n'))
    
    # Improved score calculation - more generous baseline scores
    content_quality = min(100, 75 + (15 if has_education else 0) + (10 if has_experience else 0))
    formatting_score = min(100, 75 + (word_count // 20) if word_count < 1000 else 95)
    keyword_match = min(100, 70 + (10 if has_skills else 0) + (10 if has_experience else 0) + (10 if has_education else 0))
    experience_impact = min(100, 85 if has_experience else 70)
    skills_relevance = min(100, 88 if has_skills else 75)
    ats_compatibility = min(100, 75 + (8 if has_contact else 0) + (8 if line_count > 20 else 0) + (9 if word_count > 200 else 0))
    
    overall_score = int((content_quality + formatting_score + keyword_match + experience_impact + skills_relevance + ats_compatibility) / 6)
    
    # Generate insights
    strengths = []
    if has_education:
        strengths.append("Strong educational background clearly presented")
    if has_experience:
        strengths.append("Relevant work experience highlighted effectively")
    if has_skills:
        strengths.append("Technical skills section is comprehensive")
    if word_count > 300:
        strengths.append("Adequate content length for detailed assessment")
    
    improvements = []
    if not has_education:
        improvements.append("Add education section with degree and institution details")
    if not has_experience:
        improvements.append("Include work experience with quantifiable achievements")
    if word_count < 200:
        improvements.append("Expand resume content - aim for 400-600 words")
    improvements.append("Use action verbs to describe accomplishments")
    
    missing_keywords = [
        "Leadership", "Team Collaboration", "Project Management",
        "Problem Solving", "Communication", "Agile/Scrum"
    ]
    
    suggested_skills = [
        "Cloud Technologies (AWS/Azure)", "CI/CD Pipelines",
        "Database Management", "API Development", "Version Control (Git)"
    ]
    
    return AnalyzeResponse(
        overall_score=overall_score,
        content_quality=content_quality,
        formatting_score=formatting_score,
        keyword_match=keyword_match,
        experience_impact=experience_impact,
        skills_relevance=skills_relevance,
        ats_compatibility=ats_compatibility,
        strengths=strengths[:4],
        improvements=improvements[:4],
        missing_keywords=missing_keywords[:6],
        suggested_skills=suggested_skills[:5],
        extracted_text=extracted_text[:500]
    )


@router.post("/parse-json", response_model=dict)
async def parse_resume_json(resume_data: ResumeData):
    """Parse and store structured resume JSON data"""
    
    # Calculate scores based on resume data
    total_skills = (
        len(resume_data.skills.programming_languages) +
        len(resume_data.skills.web_technologies) +
        len(resume_data.skills.tools_and_technologies) +
        len(resume_data.skills.soft_skills)
    )
    
    experience_years = len(resume_data.experience)
    education_level = len(resume_data.education)
    
    # Generate analysis
    overall_score = min(100, 70 + total_skills + (experience_years * 5) + (education_level * 3))
    
    return {
        "success": True,
        "message": "Resume parsed successfully",
        "data": resume_data.dict(),
        "analysis": {
            "overall_score": overall_score,
            "total_skills": total_skills,
            "experience_count": experience_years,
            "education_count": education_level,
            "profile_completeness": 95 if resume_data.about_me else 80
        }
    }


@router.post("/generate-better", response_model=dict)
async def generate_better_resume(file: UploadFile = File(...)):
    """Generate an improved version of the resume using AI"""
    
    try:
        # Extract text based on file type
        extracted_text = ""
        
        if file.filename.endswith('.pdf'):
            content = await file.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
                
        elif file.filename.endswith('.docx'):
            content = await file.read()
            doc = docx.Document(io.BytesIO(content))
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF or DOCX")
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the document")
        
        # Use Gemini AI to generate improved resume
        if GEMINI_API_KEY:
            try:
                prompt = f"""You are an expert resume writer and career coach. Based on the following resume, create an IMPROVED and ENHANCED version that:

1. Fixes grammar and formatting issues
2. Uses strong action verbs
3. Adds quantifiable achievements where applicable
4. Improves the professional summary
5. Optimizes for ATS (Applicant Tracking Systems)
6. Highlights key skills and accomplishments
7. Uses proper resume formatting and structure

Original Resume:
{extracted_text[:4000]}

Please generate a complete, professional, improved resume in MARKDOWN format. Include all sections: Contact Info, Professional Summary, Skills, Experience, Education, etc. Make it compelling and achievement-focused.

Return ONLY the improved resume text in markdown format, no additional explanations."""

                headers = {"Content-Type": "application/json"}
                payload = {
                    "contents": [{
                        "parts": [{"text": prompt}]
                    }]
                }
                
                response = requests.post(
                    f"{GEMINI_API_URL}?key={GEMINI_API_KEY}",
                    headers=headers,
                    json=payload,
                    timeout=45
                )
                
                if response.status_code == 200:
                    result = response.json()
                    improved_resume = result.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
                    
                    return {
                        "success": True,
                        "improved_resume": improved_resume,
                        "original_length": len(extracted_text),
                        "improved_length": len(improved_resume)
                    }
                else:
                    raise HTTPException(status_code=500, detail=f"Gemini API Error: {response.status_code}")
                    
            except Exception as ai_error:
                print(f"Gemini API Error: {ai_error}")
                raise HTTPException(status_code=500, detail=f"AI Generation Error: {str(ai_error)}")
        else:
            raise HTTPException(status_code=500, detail="Gemini API key not configured")
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating resume: {str(e)}")


@router.post("/upload", response_model=ResumeAnalysis)
async def analyze_resume(file: UploadFile = File(...)):
    """Upload and analyze resume"""
    
    if not file.filename.endswith(('.pdf', '.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    # Mock analysis - replace with actual AI analysis
    return ResumeAnalysis(
        overall_score=87,
        strengths=[
            "Strong technical skills section",
            "Quantifiable achievements",
            "Clear project descriptions",
            "Good use of action verbs"
        ],
        weaknesses=[
            "Missing keywords for ATS optimization",
            "Inconsistent formatting in experience section",
            "Summary could be more impactful"
        ],
        improvements=[
            "Add more industry-specific keywords",
            "Standardize bullet point format",
            "Include metrics in all achievements",
            "Optimize for ATS scanning"
        ],
        keyword_match=75,
        formatting_score=82,
        content_score=90,
        ats_compatibility=78,
        sections_analysis={
            "summary": {"score": 70, "suggestions": ["Make it more compelling", "Add key achievements"]},
            "experience": {"score": 85, "suggestions": ["Add more metrics"]},
            "skills": {"score": 90, "suggestions": ["Great technical coverage"]},
            "education": {"score": 80, "suggestions": ["Include relevant coursework"]}
        }
    )

@router.post("/improve", response_model=List[ResumeImprovement])
async def improve_resume_section(section: str, content: str):
    """Get AI-powered resume improvement suggestions"""
    
    # Mock improvements - replace with actual AI generation
    return [
        ResumeImprovement(
            section=section,
            current_text=content,
            improved_text="Enhanced version with better action verbs and quantifiable results",
            reasoning="Added metrics and strengthened impact statements"
        )
    ]

@router.get("/templates", response_model=List[dict])
async def get_resume_templates():
    """Get resume templates"""
    return [
        {
            "id": "modern-tech",
            "name": "Modern Tech Resume",
            "description": "Clean, ATS-friendly template for tech professionals",
            "preview_url": "/templates/modern-tech.png"
        },
        {
            "id": "executive",
            "name": "Executive Resume",
            "description": "Professional template for senior positions",
            "preview_url": "/templates/executive.png"
        }
    ]
